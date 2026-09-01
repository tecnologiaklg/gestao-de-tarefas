import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx2pdf import convert

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_callout_border(cell, color="3B82F6", sz="24"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def create_styled_document():
    doc = Document()
    
    # Page setup (A4 with 2cm margins)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(30, 41, 59) # Slate-800
    style_normal.paragraph_format.space_after = Pt(4)
    style_normal.paragraph_format.line_spacing = 1.15

    return doc

def add_inline_formatted_text(paragraph, text, default_color=None, default_size=None, default_bold=False):
    # Regex to extract formatting: links [text](url), bold **text**, italic *text*, code `code`
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            if default_color: run.font.color.rgb = default_color
            if default_size: run.font.size = default_size
        elif token.startswith('*') and token.endswith('*') and not token.startswith('**'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            if default_color: run.font.color.rgb = default_color
            if default_size: run.font.size = default_size
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(f" {token[1:-1]} ")
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(15, 23, 42)
            run.bold = True
        elif token.startswith('[') and ']' in token and '(' in token and token.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if match:
                link_text, _ = match.groups()
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(37, 99, 235)
                run.font.size = default_size or Pt(10)
                run.bold = True
        else:
            run = paragraph.add_run(token)
            run.bold = default_bold
            if default_color: run.font.color.rgb = default_color
            if default_size: run.font.size = default_size

def convert_md_to_docx(md_path, docx_path, main_title):
    doc = create_styled_document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_lang = ''
    code_lines = []
    
    in_table = False
    table_rows = []
    
    for raw_line in lines:
        line = raw_line.rstrip('\r\n')
        trimmed = line.strip()
        
        # Code block handler
        if trimmed.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                in_code_block = False
                code_lines = []
                
                # Render code block as styled callout box
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                cell = tbl.cell(0, 0)
                cell.width = Inches(6.67)
                set_cell_background(cell, "0F172A") # Slate-900
                set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(226, 232, 240) # Slate-200
                
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            else:
                in_code_block = True
                code_lang = trimmed[3:].strip()
                code_lines = []
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue
            
        # Ignore markdown anchor tags
        if re.match(r'^<a\s+id="[^"]+"><\/a>$', trimmed):
            continue
            
        # Horizontal rule
        if trimmed in ['---', '***', '===']:
            if in_table:
                in_table = False
                table_rows = []
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run('━' * 55)
            run.font.color.rgb = RGBColor(203, 213, 225)
            run.font.size = Pt(8)
            continue
            
        # Table handler
        if trimmed.startswith('|') and trimmed.endswith('|'):
            cells = [c.strip() for c in trimmed[1:-1].split('|')]
            if all(re.match(r'^:?-+:?$', c) for c in cells):
                continue # Separator row
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            # End of table, render it
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                num_rows = len(table_rows)
                tbl = doc.add_table(rows=num_rows, cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(tbl, color="CBD5E1", sz="4")
                
                col_widths = [Inches(6.67 / num_cols)] * num_cols
                
                for r_idx, row_data in enumerate(table_rows):
                    row = tbl.rows[r_idx]
                    is_header = (r_idx == 0)
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = row.cells[c_idx]
                            cell.width = col_widths[c_idx]
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                            
                            if is_header:
                                set_cell_background(cell, "F1F5F9")
                            elif r_idx % 2 == 1:
                                set_cell_background(cell, "FFFFFF")
                            else:
                                set_cell_background(cell, "F8FAFC")
                                
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(0)
                            add_inline_formatted_text(p, cell_value, default_bold=is_header)
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            table_rows = []
            in_table = False
            
        # Headings
        if trimmed.startswith('# '):
            h_text = trimmed[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            add_inline_formatted_text(p, h_text, default_color=RGBColor(15, 23, 42), default_size=Pt(16), default_bold=True)
            continue
            
        if trimmed.startswith('## '):
            h_text = trimmed[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(11)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            add_inline_formatted_text(p, h_text, default_color=RGBColor(30, 58, 138), default_size=Pt(12.5), default_bold=True)
            continue
            
        if trimmed.startswith('### '):
            h_text = trimmed[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            add_inline_formatted_text(p, h_text, default_color=RGBColor(51, 65, 85), default_size=Pt(10.5), default_bold=True)
            continue
            
        # Callouts (> [!NOTE] or > text)
        if trimmed.startswith('>'):
            callout_text = trimmed[1:].strip()
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            cell = tbl.cell(0, 0)
            cell.width = Inches(6.67)
            set_cell_background(cell, "EFF6FF") # Blue-50
            set_callout_border(cell, color="3B82F6", sz="24")
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_formatted_text(p, callout_text, default_color=RGBColor(30, 58, 138), default_size=Pt(9.5))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
            
        # Lists
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
        if list_match:
            indent_len = len(list_match.group(1))
            bullet = list_match.group(2)
            item_text = list_match.group(3)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 if indent_len == 0 else 0.45)
            p.paragraph_format.space_after = Pt(2)
            
            bullet_run = p.add_run("▪  " if bullet in ['-', '*'] else f"{bullet} ")
            bullet_run.bold = True
            bullet_run.font.color.rgb = RGBColor(59, 130, 246)
            
            add_inline_formatted_text(p, item_text)
            continue
            
        # Regular text
        if trimmed:
            p = doc.add_paragraph()
            add_inline_formatted_text(p, trimmed)
            
    doc.save(docx_path)
    print(f"[OK] {docx_path} gerado com sucesso!")

def main():
    root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    docs = [
        ('manual_colaborador.md', 'manual_colaborador.docx', 'manual_colaborador.pdf', 'Manual do Usuário - Portal de Gestão de Tarefas'),
        ('manual_root.md', 'manual_root.docx', 'manual_root.pdf', 'Manual do Administrador (Root) - Portal de Gestão de Tarefas'),
    ]
    
    for md_file, docx_file, pdf_file, title in docs:
        md_path = os.path.join(root_dir, md_file)
        docx_path = os.path.join(root_dir, docx_file)
        pdf_path = os.path.join(root_dir, pdf_file)
        
        print(f"\n--- Processando {docx_file} ---")
        convert_md_to_docx(md_path, docx_path, title)
        
        print(f"Convertendo {docx_file} -> {pdf_file} via Microsoft Word...")
        convert(docx_path, pdf_path)
        print(f"[OK] {pdf_file} gerado com sucesso a partir do DOCX!")

if __name__ == '__main__':
    main()
